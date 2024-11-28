# Importations standard
import os
from io import BytesIO

# Importations Django
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.conf import settings
from django.core.files.storage import default_storage

# Importations tierces
import openai
import pdfplumber
from docx import Document
from reportlab.pdfgen import canvas

# Importations des modèles et formulaires
from .models import User, JobOffer, Candidat, Recruteur, Candidature
from .forms import JobForm, RegistrationForm, LoginForm, CandidatForm, CandidatureForm

# Importations des utilitaires personnalisés
from rcw.cv_processing import extract_competences
from .scraper import scrape_indeed_jobs

# Configuration OpenAI
openai.api_key = settings.OPENAI_API_KEY


# ---- Fonctions Utilitaires ----

def extract_text_from_cv(cv_file):
    """Extrait le texte d'un fichier CV (PDF ou Word)."""
    file_extension = os.path.splitext(cv_file.name)[1].lower()
    try:
        if file_extension == '.pdf':
            text = ""
            with pdfplumber.open(cv_file) as pdf:
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
            return text.strip()
        elif file_extension in ['.doc', '.docx']:
            doc = Document(cv_file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()
        else:
            raise ValueError("Format non pris en charge. Utilisez un fichier PDF ou Word.")
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte : {e}")
        return ""

def calculer_pertinence_ia(offre, cv_texte, lettre_motivation_texte):
    """Calcule la pertinence d'un candidat pour une offre d'emploi via OpenAI GPT."""
    prompt = f"""
    Évaluez la pertinence d'un candidat pour une offre d'emploi.

    Offre d'emploi :
    - Titre : {offre.titre}
    - Description : {offre.description}
    - Compétences requises : {offre.competences_requises}
    - Localisation : {offre.localisation}

    CV :
    {cv_texte}

    Lettre de motivation :
    {lettre_motivation_texte}

    Donnez une note de pertinence stricte entre 0 et 100.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300
        )
        score = response.choices[0].message['content'].strip()
        return int(score) if score.isdigit() else 0
    except Exception as e:
        print(f"Erreur lors du calcul de pertinence : {e}")
        return 0


# ---- Vues Principales ----

def index(request):
    """Affiche les offres d'emploi récentes sur la page d'accueil."""
    jobs = JobOffer.objects.order_by('-date_publication')[:6]
    return render(request, 'index.html', {'jobs': jobs})

@login_required
def profile(request):
    """Affiche le profil de l'utilisateur connecté."""
    user = request.user
    context = {'user': user}

    if user.role == 'RECRUTEUR':
        try:
            recruteur = Recruteur.objects.get(user=user)
            offres = JobOffer.objects.filter(recruteur=recruteur)
            context['offres'] = offres
        except Recruteur.DoesNotExist:
            context['offres'] = []

    return render(request, 'profile.html', context)

def register(request):
    """Permet l'inscription d'un nouvel utilisateur."""
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            try:
                user = form.save()
                login(request, user)
                return redirect('index')
            except Exception as e:
                form.add_error(None, f"Erreur : {str(e)}")
    else:
        form = RegistrationForm()
    return render(request, 'register.html', {'form': form})

def login_view(request):
    """Connexion utilisateur."""
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            user = authenticate(
                request, 
                username=form.cleaned_data['username'], 
                password=form.cleaned_data['password']
            )
            if user:
                login(request, user)
                return redirect('index')
            else:
                form.add_error(None, 'Nom d’utilisateur ou mot de passe incorrect')
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})

@login_required
def logout_view(request):
    """Déconnexion utilisateur."""
    logout(request)
    return redirect('index')


# ---- Gestion des Offres d'Emploi ----

@login_required
def post_job(request):
    """Permet à un recruteur de publier une offre d'emploi."""
    if request.user.role == 'RECRUTEUR':
        if not hasattr(request.user, 'recruteur_profile'):
            Recruteur.objects.create(user=request.user)

        if request.method == 'POST':
            form = JobForm(request.POST)
            if form.is_valid():
                job = form.save(commit=False)
                job.recruteur = request.user.recruteur_profile
                job.save()
                return redirect('job_list')
        else:
            form = JobForm()
        return render(request, 'post_job.html', {'form': form})
    return redirect('job_list')

@login_required
def job_list(request):
    """Affiche la liste des offres d'emploi."""
    jobs = JobOffer.objects.all()
    return render(request, 'job_list.html', {'jobs': jobs})


# ---- Génération et Téléchargement ----

@login_required
def generate_cover_letter(request, offre_id):
    """Génère une lettre de motivation via OpenAI."""
    offre = get_object_or_404(JobOffer, id=offre_id)
    candidat = request.user.candidat_profile
    prompt = f"Rédige une lettre de motivation pour l'offre '{offre.titre}'..."

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500
        )
        letter = response.choices[0].message['content']
        return render(request, "cover_letter.html", {"offre": offre, "letter": letter})
    except Exception as e:
        print(f"Erreur OpenAI : {e}")
        return HttpResponse("Une erreur est survenue.")

@login_required
def download_cover_letter_pdf(request, offre_id):
    """Télécharge une lettre de motivation au format PDF."""
    # Logique pour générer et renvoyer un PDF.
    pass


# ---- Chatbot ----

def chatbot_response(request):
    """Génère une réponse via OpenAI pour un chatbot intégré."""
    if request.method == 'POST':
        user_message = request.POST.get('message')
        # Logique pour générer une réponse.
    return JsonResponse({"error": "Méthode non autorisée"}, status=405)

# ---- Gestion des Candidatures ----

@login_required
def postuler_offre(request, offre_id):
    """Permet à un candidat de postuler à une offre d'emploi."""
    offre = get_object_or_404(JobOffer, id=offre_id)
    allowed_extensions = ['pdf', 'doc', 'docx']

    if request.method == 'POST':
        cv = request.FILES.get('cv')
        lettre_motivation = request.FILES.get('lettre_motivation')

        if cv and lettre_motivation:
            cv_extension = os.path.splitext(cv.name)[1][1:].lower()
            lettre_extension = os.path.splitext(lettre_motivation.name)[1][1:].lower()

            if cv_extension in allowed_extensions and lettre_extension in allowed_extensions:
                candidature = Candidature(
                    candidat=request.user,
                    offre=offre,
                    cv=cv,
                    lettre_motivation=lettre_motivation
                )
                candidature.save()
                messages.success(request, "Votre candidature a été soumise avec succès.")
                return redirect('job_list')
            else:
                messages.error(request, "Veuillez télécharger des fichiers au format PDF, DOC ou DOCX.")
        else:
            messages.error(request, "Veuillez télécharger un CV et une lettre de motivation.")

    return render(request, 'postuler_offre.html', {'offre': offre})

@login_required
def mes_candidatures(request):
    """Affiche les candidatures envoyées par le candidat."""
    candidatures = Candidature.objects.filter(candidat=request.user)
    return render(request, 'mes_candidatures.html', {'candidatures': candidatures})


# ---- Recherche et Analyse ----

def search_job_offers(request):
    """Recherche des offres d'emploi en fonction d'une requête utilisateur."""
    query = request.GET.get('q')
    job_offers = JobOffer.objects.filter(
        titre__icontains=query
    ) | JobOffer.objects.filter(
        description__icontains=query
    ) if query else JobOffer.objects.all()

    return render(request, 'search_job_offers.html', {
        'job_offers': job_offers,
        'query': query
    })

def match_jobs_view(request):
    """Associe un CV téléchargé à des offres d'emploi pertinentes."""
    if request.method == "POST" and request.FILES.get("cv_file"):
        cv_file = request.FILES["cv_file"]
        try:
            cv_text = extract_text_from_cv(cv_file)
            if not cv_text.strip():
                return render(request, "cv_analysis_result.html", {"message": "Le fichier CV est vide ou non lisible."})

            competences = extract_competences(cv_text)
            matching_jobs = JobOffer.objects.none()

            for competence in competences:
                matching_jobs |= JobOffer.objects.filter(competences_requises__icontains=competence)

            return render(request, "cv_analysis_result.html", {
                "jobs": matching_jobs if matching_jobs.exists() else None,
                "message": "Aucune correspondance trouvée." if not matching_jobs.exists() else None
            })
        except Exception as e:
            print("Erreur de lecture du fichier :", e)
            return render(request, "cv_analysis_result.html", {"message": "Impossible de lire le fichier CV."})

    return render(request, "index.html")


# ---- Gestion des Offres d'Emploi par le Recruteur ----

@login_required
def edit_job(request, job_id):
    """Permet à un recruteur de modifier une offre d'emploi existante."""
    job = get_object_or_404(JobOffer, id=job_id)

    if request.user.role == 'RECRUTEUR' and job.recruteur.user == request.user:
        if request.method == 'POST':
            form = JobForm(request.POST, instance=job)
            if form.is_valid():
                form.save()
                return redirect('job_list')
        else:
            form = JobForm(instance=job)
        return render(request, 'edit_job.html', {'form': form, 'job': job})
    return redirect('job_list')

@login_required
def delete_job(request, job_id):
    """Permet à un recruteur de supprimer une offre d'emploi."""
    job = get_object_or_404(JobOffer, id=job_id)

    if request.user.role == 'RECRUTEUR' and job.recruteur.user == request.user:
        if request.method == 'POST':
            job.delete()
            return redirect('job_list')
        return render(request, 'confirm_delete.html', {'job': job})
    return redirect('job_list')


# ---- Chatbot ----

def chatbot_page(request):
    """Affiche la page du chatbot."""
    return render(request, "chatbot.html")

def chatbot_response(request):
    """Génère une réponse via le chatbot intégré en fonction de la question de l'utilisateur."""
    if request.method == 'POST':
        user_message = request.POST.get('message')

        if not user_message.strip():
            return JsonResponse({"error": "Le message ne peut pas être vide."}, status=400)

        try:
            offers_text = "\n".join([
                f"{job.titre}, {job.localisation}, {job.salaire}$"
                for job in JobOffer.objects.all()
            ])

            prompt = (
                f"Vous êtes un assistant virtuel pour un site d'emploi. "
                f"Voici les offres disponibles : {offers_text}. "
                f"L'utilisateur a demandé : {user_message}."
            )

            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500
            )
            chatbot_response = response.choices[0].message['content']
            return JsonResponse({"response": chatbot_response})

        except Exception as e:
            print(f"Erreur OpenAI : {e}")
            return JsonResponse({"error": "Erreur du chatbot."}, status=500)

    return JsonResponse({"error": "Méthode non autorisée."}, status=405)


# ---- Détails et Pages Statique ----

def job_detail(request, offre_id):
    """Affiche les détails d'une offre d'emploi."""
    offre = get_object_or_404(JobOffer, id=offre_id)
    return render(request, 'job_detail.html', {'offre': offre})

def about_view(request):
    """Affiche la page À propos."""
    sections = [
        {'title': 'Accessibilité', 'text': 'Notre mission est de rendre les emplois accessibles.'}
    ]
    return render(request, 'about.html', {'sections': sections})

# téléchargerment des cv

@login_required
def upload_resume(request):
    """Permet aux candidats de télécharger leur CV et lettre de motivation."""
    if request.method == 'POST':
        form = CandidatForm(request.POST, request.FILES)
        if form.is_valid():
            candidat_profile, created = Candidat.objects.get_or_create(user=request.user)
            if form.cleaned_data['cv']:
                candidat_profile.cv = form.cleaned_data['cv']
            if form.cleaned_data['lettres_motivation']:
                candidat_profile.lettres_motivation = form.cleaned_data['lettres_motivation']
            candidat_profile.competences = form.cleaned_data['competences']
            candidat_profile.save()
            messages.success(request, "Vos documents ont été téléchargés avec succès.")
            return redirect('profile')
    else:
        form = CandidatForm()
    return render(request, 'upload_resume.html', {'form': form})


#Affichage des candidatures d’un recruteur

@login_required
def candidatures_par_offre(request, offre_id):
    """Affiche toutes les candidatures reçues pour une offre spécifique."""
    offre = get_object_or_404(JobOffer, id=offre_id)
    candidatures = Candidature.objects.filter(offre=offre)

    # Calcul de pertinence pour chaque candidature
    for candidature in candidatures:
        try:
            cv_text = extract_text_from_cv(candidature.cv)
            lettre_text = extract_text_from_cv(candidature.lettre_motivation)
            candidature.pertinence = calculer_pertinence_ia(offre, cv_text, lettre_text)
        except Exception as e:
            print(f"Erreur pour la candidature {candidature.id} : {e}")
            candidature.pertinence = 0

    # Trier les candidatures par pertinence décroissante
    candidatures = sorted(candidatures, key=lambda c: c.pertinence, reverse=True)

    return render(request, 'candidatures_par_offre.html', {
        'offre': offre,
        'candidatures': candidatures
    })

#Suppression de candidatures par un recruteur

@login_required
def supprimer_candidature(request, candidature_id):
    """Permet de supprimer une candidature spécifique."""
    candidature = get_object_or_404(Candidature, id=candidature_id)

    if request.method == 'POST' and request.user.role == 'RECRUTEUR':
        candidature.delete()
        messages.success(request, "Candidature supprimée avec succès.")
        return redirect('mes_candidatures')

    return render(request, 'confirm_delete.html', {'candidature': candidature})


#Extraction du texte pour CV et lettres

def extract_text_from_cv(cv_file):
    """Extrait le texte d'un fichier CV (PDF ou Word)."""
    file_extension = os.path.splitext(cv_file.name)[1].lower()
    try:
        if file_extension == '.pdf':
            text = ""
            with pdfplumber.open(cv_file) as pdf:
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
            return text.strip()
        elif file_extension in ['.doc', '.docx']:
            doc = Document(cv_file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()
        else:
            raise ValueError("Format non pris en charge. Utilisez un fichier PDF ou Word.")
    except Exception as e:
        print(f"Erreur lors de l'extraction : {e}")
        return ""
    
    
#Gestion de l’édition et suppression des offres

@login_required
def edit_job(request, job_id):
    """Permet de modifier une offre d'emploi existante."""
    job = get_object_or_404(JobOffer, id=job_id)

    if request.user.role == 'RECRUTEUR' and job.recruteur.user == request.user:
        if request.method == 'POST':
            form = JobForm(request.POST, instance=job)
            if form.is_valid():
                form.save()
                messages.success(request, "Offre mise à jour avec succès.")
                return redirect('job_list')
        else:
            form = JobForm(instance=job)
        return render(request, 'edit_job.html', {'form': form, 'job': job})

    return redirect('job_list')


#supprimer une offre 

@login_required
def delete_job(request, job_id):
    """Supprime une offre d'emploi."""
    job = get_object_or_404(JobOffer, id=job_id)

    if request.user.role == 'RECRUTEUR' and job.recruteur.user == request.user:
        if request.method == 'POST':
            job.delete()
            messages.success(request, "Offre supprimée avec succès.")
            return redirect('job_list')

    return render(request, 'confirm_delete.html', {'job': job})



